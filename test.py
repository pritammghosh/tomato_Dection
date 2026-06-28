"""
Professional tomato grading report generator.

The script runs a YOLO segmentation model on sample tomato images and creates
an A4 DOCX report with an executive summary, image-level dashboards, and
per-tomato defect details.
"""

from __future__ import annotations

import argparse
import importlib.util
import io
import random
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path


BASE_DIR = Path(__file__).resolve().parent

# =====================================================
# REPORT SETTINGS - edit these values before running
# =====================================================

# Number of images to test in each report.
SAMPLE_COUNT = 10

# True = choose random images. False = use the first SAMPLE_COUNT images in filename order.
USE_RANDOM_SELECTION = True

# Change this number to get a different random selection.
# Keep it the same when you want repeatable results.
RANDOM_SEED = 100


REQUIRED_PACKAGES = {
    "docx": "python-docx",
    "matplotlib": "matplotlib",
    "numpy": "numpy",
    "PIL": "Pillow",
    "ultralytics": "ultralytics",
}


def use_project_virtual_environment() -> None:
    project_python = BASE_DIR / ".venv" / "Scripts" / "python.exe"
    current_python = Path(sys.executable).resolve()

    if project_python.exists() and current_python != project_python.resolve():
        completed = subprocess.run(
            [str(project_python), str(BASE_DIR / Path(__file__).name), *sys.argv[1:]],
            cwd=BASE_DIR,
            check=False,
        )
        raise SystemExit(completed.returncode)


def ensure_dependencies() -> None:
    missing = [
        package_name
        for import_name, package_name in REQUIRED_PACKAGES.items()
        if importlib.util.find_spec(import_name) is None
    ]
    if not missing:
        return

    install_command = f'"{sys.executable}" -m pip install {" ".join(missing)}'
    message = (
        "Missing required Python package(s): "
        f"{', '.join(missing)}\n\n"
        "Fix options:\n"
        f"1. Run this script with the project virtual environment:\n"
        f"   .\\.venv\\Scripts\\python.exe test.py --samples 10\n\n"
        f"2. Or install the missing package(s) into the current Python:\n"
        f"   {install_command}"
    )
    raise SystemExit(message)


use_project_virtual_environment()
ensure_dependencies()

import docx  # noqa: E402
import matplotlib.pyplot as plt  # noqa: E402
import numpy as np  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Mm, Pt, RGBColor  # noqa: E402
from PIL import Image  # noqa: E402
from ultralytics import YOLO  # noqa: E402


CLASS_DEFECTIVE = 0
CLASS_RIPE = 1
CLASS_UNRIPE = 2

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
DEFAULT_MODEL_PATH = BASE_DIR / "best.pt"
DEFAULT_IMAGE_DIR = BASE_DIR / "images"
DEFAULT_OUTPUT_DIR = BASE_DIR / "results"

DEFECT_THRESHOLD_PERCENT = 5.0
MIN_OVERLAP_PIXELS = 20

COLORS = {
    "navy": "#152238",
    "blue": "#2563EB",
    "muted": "#667085",
    "line": "#D0D5DD",
    "ripe": "#2E7D32",
    "unripe": "#F9A825",
    "defective": "#C62828",
    "background": "#F8FAFC",
    "panel": "#EEF4FF",
    "dark_panel": "#101828",
    "white": "#FFFFFF",
}


@dataclass
class Detection:
    mask: np.ndarray
    class_id: int
    confidence: float
    box: np.ndarray


@dataclass
class TomatoAssessment:
    number: int
    label: str
    status: str
    color: str
    defect_percent: float
    confidence: float
    center_x: int
    center_y: int
    mask: np.ndarray
    defect_mask: np.ndarray
    box: np.ndarray


@dataclass
class ImageSummary:
    image_name: str
    tomato_count: int
    ripe_count: int
    unripe_count: int
    defective_count: int
    ripe_pct: float
    unripe_pct: float
    defect_pct: float


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a professional tomato grading DOCX report.")
    parser.add_argument("--model", default=DEFAULT_MODEL_PATH, type=Path, help="Path to YOLO model weights.")
    parser.add_argument("--images", default=DEFAULT_IMAGE_DIR, type=Path, help="Folder containing images to analyze.")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_DIR, type=Path, help="Folder for generated reports.")
    parser.add_argument("--samples", default=SAMPLE_COUNT, type=int, help="Number of images to include.")
    parser.add_argument("--confidence", default=0.25, type=float, help="YOLO confidence threshold.")
    parser.add_argument("--seed", default=RANDOM_SEED, type=int, help="Random seed for repeatable image sampling.")
    parser.add_argument(
        "--random",
        dest="use_random_selection",
        action="store_true",
        default=USE_RANDOM_SELECTION,
        help="Choose random images for the report.",
    )
    parser.add_argument(
        "--no-random",
        dest="use_random_selection",
        action="store_false",
        help="Use the first images in filename order instead of random images.",
    )
    return parser.parse_args()


def configure_plotting() -> None:
    plt.style.use("seaborn-v0_8-whitegrid")
    plt.rcParams.update(
        {
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": COLORS["line"],
            "axes.titleweight": "bold",
            "font.size": 10,
        }
    )


def validate_inputs(model_path: Path, image_dir: Path, sample_count: int) -> None:
    if not model_path.exists():
        raise FileNotFoundError(f"Model file not found: {model_path.resolve()}")
    if not image_dir.exists():
        raise FileNotFoundError(f"Image folder not found: {image_dir.resolve()}")
    if sample_count <= 0:
        raise ValueError("--samples must be greater than zero.")


def collect_images(image_dir: Path) -> list[Path]:
    images = sorted(path for path in image_dir.iterdir() if path.suffix.lower() in IMAGE_EXTENSIONS)
    if not images:
        raise FileNotFoundError(f"No supported image files found in: {image_dir.resolve()}")
    return images


def select_images(image_paths: list[Path], sample_count: int, use_random_selection: bool, seed: int) -> list[Path]:
    selected_count = min(sample_count, len(image_paths))
    if use_random_selection:
        random_generator = random.Random(seed)
        return random_generator.sample(image_paths, selected_count)
    return image_paths[:selected_count]


def make_report_path(output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return output_dir / f"Tomato_Grading_Report_{timestamp}.docx"


def setup_document() -> docx.Document:
    document = docx.Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(52, 64, 84)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(31, 42, 68)
        style.font.bold = True

    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "Automated Tomato Grading | Quality Assurance Report"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(102, 112, 133)

        footer = section.footer.paragraphs[0]
        footer.text = "Generated by AI segmentation pipeline"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(102, 112, 133)

    return document


def add_heading(document: docx.Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(6)


def hex_to_rgb(color: str) -> RGBColor:
    color = color.lstrip("#")
    return RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color: str = "D0D5DD", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.lstrip("#"))


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = COLORS["navy"],
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = hex_to_rgb(color)


def add_metric_table(document: docx.Document, rows: list[tuple[str, str]], accent: str = COLORS["panel"]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], accent)
        set_cell_shading(cells[1], COLORS["white"])
        set_cell_border(cells[0])
        set_cell_border(cells[1])
        set_cell_text(cells[0], label.upper(), bold=True, color=COLORS["muted"], size=8)
        set_cell_text(cells[1], value, bold=True, color=COLORS["navy"], size=10)


def add_labeled_paragraph(document: docx.Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = hex_to_rgb(COLORS["navy"])
    text_run = paragraph.add_run(text)
    text_run.font.color.rgb = hex_to_rgb(COLORS["muted"])


def add_figure_to_doc(document: docx.Document, figure: plt.Figure, width_inches: float = 7.2) -> None:
    buffer = io.BytesIO()
    figure.savefig(buffer, format="png", bbox_inches="tight", dpi=180)
    buffer.seek(0)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(buffer, width=Inches(width_inches))


def add_title_page(document: docx.Document, sample_count: int, model_path: Path, image_dir: Path) -> None:
    cover = document.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover.rows[0].cells[0]
    set_cell_shading(cover_cell, COLORS["dark_panel"])
    set_cell_border(cover_cell, COLORS["dark_panel"])

    paragraph = cover_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(22)
    paragraph.paragraph_format.space_after = Pt(8)
    eyebrow = paragraph.add_run("QUALITY ASSURANCE REPORT")
    eyebrow.bold = True
    eyebrow.font.size = Pt(10)
    eyebrow.font.color.rgb = hex_to_rgb("#93C5FD")

    title = cover_cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Automated Tomato\nGrading Report")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = hex_to_rgb(COLORS["white"])

    subtitle = cover_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("AI segmentation, ripeness classification, and defect analysis")
    run.font.size = Pt(12)
    run.font.color.rgb = hex_to_rgb("#D0D5DD")

    document.add_paragraph()
    add_metric_table(
        document,
        [
            ("Generated on", datetime.now().strftime("%d %b %Y, %I:%M %p")),
            ("Images selected", str(sample_count)),
            ("Model", model_path.name),
            ("Image source", str(image_dir)),
            ("Defect threshold", f"{DEFECT_THRESHOLD_PERCENT:.1f}% of tomato mask"),
        ],
        accent="#E0EAFF",
    )
    document.add_paragraph()
    add_labeled_paragraph(
        document,
        "Purpose",
        "This report summarizes tomato quality from model-predicted segmentation masks and defect regions.",
    )
    add_labeled_paragraph(
        document,
        "Classification rule",
        "A tomato is marked defective when overlapping defect area exceeds the configured threshold.",
    )
    document.add_page_break()


def split_detections(result) -> tuple[list[Detection], list[Detection], tuple[int, int]]:
    tomatoes: list[Detection] = []
    defects: list[Detection] = []

    if result.masks is None or result.boxes is None or len(result.masks.data) == 0:
        image_shape = getattr(result, "orig_shape", (10, 10))
        return tomatoes, defects, tuple(image_shape[:2])

    masks = result.masks.data.cpu().numpy().astype(bool)
    classes = result.boxes.cls.cpu().numpy()
    confidences = result.boxes.conf.cpu().numpy()
    boxes = result.boxes.xyxy.cpu().numpy()
    mask_shape = tuple(masks[0].shape)

    for mask, class_id, confidence, box in zip(masks, classes, confidences, boxes):
        detection = Detection(mask=mask, class_id=int(class_id), confidence=float(confidence), box=box)
        if detection.class_id in {CLASS_RIPE, CLASS_UNRIPE}:
            tomatoes.append(detection)
        elif detection.class_id == CLASS_DEFECTIVE:
            defects.append(detection)

    return tomatoes, defects, mask_shape


def calculate_quality_percentages(
    tomatoes: list[Detection], defects: list[Detection], mask_shape: tuple[int, int]
) -> tuple[float, float, float]:
    ripe_mask = np.zeros(mask_shape, dtype=bool)
    unripe_mask = np.zeros(mask_shape, dtype=bool)
    defect_mask = np.zeros(mask_shape, dtype=bool)

    for tomato in tomatoes:
        if tomato.class_id == CLASS_RIPE:
            ripe_mask |= tomato.mask
        else:
            unripe_mask |= tomato.mask

    for defect in defects:
        defect_mask |= defect.mask

    ripe_area = int(np.sum(ripe_mask & ~defect_mask))
    unripe_area = int(np.sum(unripe_mask & ~defect_mask))
    defect_area = int(np.sum(defect_mask))
    total_area = ripe_area + unripe_area + defect_area

    if total_area == 0:
        return 0.0, 0.0, 0.0

    return (
        ripe_area / total_area * 100,
        unripe_area / total_area * 100,
        defect_area / total_area * 100,
    )


def assess_tomatoes(tomatoes: list[Detection], defects: list[Detection]) -> list[TomatoAssessment]:
    assessments: list[TomatoAssessment] = []

    for index, tomato in enumerate(tomatoes, start=1):
        tomato_area = int(np.sum(tomato.mask))
        defect_union = np.zeros_like(tomato.mask, dtype=bool)

        for defect in defects:
            overlap = tomato.mask & defect.mask
            if int(np.sum(overlap)) > MIN_OVERLAP_PIXELS:
                defect_union |= overlap

        defect_area = int(np.sum(defect_union))
        defect_percent = (defect_area / tomato_area * 100) if tomato_area else 0.0

        if defect_percent > DEFECT_THRESHOLD_PERCENT:
            status = "Defective"
            label = f"Defective ({defect_percent:.1f}%)"
            color = COLORS["defective"]
        elif tomato.class_id == CLASS_RIPE:
            status = "Ripe"
            label = "Ripe"
            color = COLORS["ripe"]
        else:
            status = "Unripe"
            label = "Unripe"
            color = COLORS["unripe"]

        y_indices, x_indices = np.where(tomato.mask)
        if len(x_indices) == 0:
            continue

        assessments.append(
            TomatoAssessment(
                number=index,
                label=label,
                status=status,
                color=color,
                defect_percent=defect_percent,
                confidence=tomato.confidence,
                center_x=int(np.mean(x_indices)),
                center_y=int(np.mean(y_indices)),
                mask=tomato.mask,
                defect_mask=defect_union,
                box=tomato.box,
            )
        )

    return assessments


def assess_standalone_defects(defects: list[Detection], tomatoes: list[Detection], start_number: int) -> list[TomatoAssessment]:
    assessments: list[TomatoAssessment] = []
    number = start_number

    for defect in defects:
        has_parent_tomato = any(int(np.sum(tomato.mask & defect.mask)) > MIN_OVERLAP_PIXELS for tomato in tomatoes)
        if has_parent_tomato:
            continue

        y_indices, x_indices = np.where(defect.mask)
        if len(x_indices) == 0:
            continue

        number += 1
        assessments.append(
            TomatoAssessment(
                number=number,
                label="Defective (100%)",
                status="Defective",
                color=COLORS["defective"],
                defect_percent=100.0,
                confidence=defect.confidence,
                center_x=int(np.mean(x_indices)),
                center_y=int(np.mean(y_indices)),
                mask=defect.mask,
                defect_mask=defect.mask,
                box=defect.box,
            )
        )

    return assessments


def safe_crop_bounds(box: np.ndarray, image_shape: tuple[int, ...], padding: int = 20) -> tuple[int, int, int, int]:
    height, width = image_shape[:2]
    x1, y1, x2, y2 = [int(round(value)) for value in box]
    return (
        max(0, x1 - padding),
        max(0, y1 - padding),
        min(width, x2 + padding),
        min(height, y2 + padding),
    )


def mask_crop_bounds(mask: np.ndarray, padding: int = 20) -> tuple[int, int, int, int] | None:
    y_indices, x_indices = np.where(mask)
    if len(x_indices) == 0:
        return None
    return (
        max(0, int(x_indices.min()) - padding),
        max(0, int(y_indices.min()) - padding),
        min(mask.shape[1], int(x_indices.max()) + padding),
        min(mask.shape[0], int(y_indices.max()) + padding),
    )


def create_detail_figure(assessment: TomatoAssessment, image_array: np.ndarray) -> plt.Figure | None:
    image_x1, image_y1, image_x2, image_y2 = safe_crop_bounds(assessment.box, image_array.shape)
    crop_image = image_array[image_y1:image_y2, image_x1:image_x2]

    mask_bounds = mask_crop_bounds(assessment.mask)
    if mask_bounds is None or crop_image.size == 0:
        return None

    mask_x1, mask_y1, mask_x2, mask_y2 = mask_bounds
    crop_mask = assessment.mask[mask_y1:mask_y2, mask_x1:mask_x2]
    crop_defect = assessment.defect_mask[mask_y1:mask_y2, mask_x1:mask_x2]

    figure, axes = plt.subplots(1, 3, figsize=(14, 4.8))
    figure.patch.set_facecolor("white")
    figure.suptitle(
        f"Tomato {assessment.number}: {assessment.label}",
        fontsize=15,
        fontweight="bold",
        color=assessment.color,
    )

    axes[0].imshow(crop_image)
    axes[0].set_title("Cropped Image")
    axes[0].axis("off")

    axes[1].imshow(crop_mask, cmap="gray")
    axes[1].set_title("Tomato Mask")
    axes[1].axis("off")

    axes[2].imshow(crop_mask, cmap="gray")
    overlay = np.zeros((crop_mask.shape[0], crop_mask.shape[1], 4))
    overlay[crop_defect] = [0.78, 0.16, 0.16, 0.85]
    axes[2].imshow(overlay)
    axes[2].set_title(f"Defect Area: {assessment.defect_percent:.1f}%")
    axes[2].axis("off")

    for axis in axes:
        axis.set_facecolor(COLORS["background"])

    figure.tight_layout(rect=(0, 0, 1, 0.9))
    return figure


def add_stacked_quality_bar(axis, ripe_pct: float, unripe_pct: float, defect_pct: float) -> None:
    axis.barh([""], [ripe_pct], color=COLORS["ripe"], edgecolor="white", height=0.45)
    axis.barh([""], [unripe_pct], left=ripe_pct, color=COLORS["unripe"], edgecolor="white", height=0.45)
    axis.barh([""], [defect_pct], left=ripe_pct + unripe_pct, color=COLORS["defective"], edgecolor="white", height=0.45)

    segments = [
        ("Ripe", ripe_pct, 0, "white"),
        ("Unripe", unripe_pct, ripe_pct, "black"),
        ("Defective", defect_pct, ripe_pct + unripe_pct, "white"),
    ]
    for label, value, left, text_color in segments:
        if value >= 5:
            axis.text(
                left + value / 2,
                0,
                f"{label} {value:.1f}%",
                ha="center",
                va="center",
                color=text_color,
                fontweight="bold",
                fontsize=11,
            )

    axis.set_xlim(0, 100)
    axis.set_title("Overall Quality Composition", pad=8)
    axis.set_xticks([])
    axis.set_yticks([])
    for spine in axis.spines.values():
        spine.set_visible(False)


def create_summary_figure(summaries: list[ImageSummary]) -> plt.Figure:
    total_images = len(summaries)
    total_regions = sum(summary.tomato_count for summary in summaries)
    total_ripe = sum(summary.ripe_count for summary in summaries)
    total_unripe = sum(summary.unripe_count for summary in summaries)
    total_defective = sum(summary.defective_count for summary in summaries)
    avg_ripe = np.mean([summary.ripe_pct for summary in summaries]) if summaries else 0.0
    avg_unripe = np.mean([summary.unripe_pct for summary in summaries]) if summaries else 0.0
    avg_defect = np.mean([summary.defect_pct for summary in summaries]) if summaries else 0.0

    figure = plt.figure(figsize=(14, 6), constrained_layout=True)
    figure.patch.set_facecolor(COLORS["background"])
    grid = figure.add_gridspec(2, 4, height_ratios=[1, 1.25])

    cards = [
        ("Images", total_images, COLORS["blue"]),
        ("Regions", total_regions, COLORS["navy"]),
        ("Ripe", total_ripe, COLORS["ripe"]),
        ("Defective", total_defective, COLORS["defective"]),
    ]

    for index, (label, value, color) in enumerate(cards):
        axis = figure.add_subplot(grid[0, index])
        axis.set_facecolor("white")
        axis.text(0.06, 0.72, label.upper(), transform=axis.transAxes, fontsize=9, fontweight="bold", color=COLORS["muted"])
        axis.text(0.06, 0.26, str(value), transform=axis.transAxes, fontsize=28, fontweight="bold", color=color)
        axis.set_xticks([])
        axis.set_yticks([])
        for spine in axis.spines.values():
            spine.set_color(COLORS["line"])

    ratio_axis = figure.add_subplot(grid[1, :2])
    add_stacked_quality_bar(ratio_axis, avg_ripe, avg_unripe, avg_defect)
    ratio_axis.set_title("Average Area Composition", fontsize=12, fontweight="bold", color=COLORS["navy"])

    count_axis = figure.add_subplot(grid[1, 2:])
    count_labels = ["Ripe", "Unripe", "Defective"]
    count_values = [total_ripe, total_unripe, total_defective]
    count_colors = [COLORS["ripe"], COLORS["unripe"], COLORS["defective"]]
    count_axis.bar(count_labels, count_values, color=count_colors, width=0.55)
    count_axis.set_title("Classification Count", fontsize=12, fontweight="bold", color=COLORS["navy"])
    count_axis.grid(axis="y", alpha=0.18)
    count_axis.spines["top"].set_visible(False)
    count_axis.spines["right"].set_visible(False)
    count_axis.spines["left"].set_color(COLORS["line"])
    count_axis.spines["bottom"].set_color(COLORS["line"])
    for index, value in enumerate(count_values):
        count_axis.text(index, value + max(count_values + [1]) * 0.03, str(value), ha="center", fontweight="bold")

    return figure


def create_dashboard_figure(
    image_path: Path,
    original_image: Image.Image,
    plotted_prediction: np.ndarray,
    assessments: list[TomatoAssessment],
    quality: tuple[float, float, float],
    mask_shape: tuple[int, int],
) -> plt.Figure:
    ripe_pct, unripe_pct, defect_pct = quality
    figure = plt.figure(figsize=(15, 9), constrained_layout=True)
    figure.patch.set_facecolor(COLORS["background"])
    figure.suptitle(f"Image Analysis: {image_path.name}", fontsize=16, fontweight="bold", color=COLORS["navy"])
    grid = figure.add_gridspec(3, 2, height_ratios=[4.5, 4.5, 1], hspace=0.32, wspace=0.06)

    original_axis = figure.add_subplot(grid[0:2, 0])
    original_axis.imshow(original_image)
    original_axis.set_title("Original Image")
    original_axis.axis("off")
    original_axis.set_facecolor("white")

    prediction_axis = figure.add_subplot(grid[0:2, 1])
    prediction_axis.imshow(plotted_prediction)
    prediction_axis.set_title("YOLO Segmentation and Classification")
    prediction_axis.axis("off")
    prediction_axis.set_facecolor("white")

    plot_height, plot_width = plotted_prediction.shape[:2]
    mask_height, mask_width = mask_shape
    x_scale = plot_width / mask_width if mask_width else 1
    y_scale = plot_height / mask_height if mask_height else 1

    for assessment in assessments:
        prediction_axis.text(
            assessment.center_x * x_scale,
            assessment.center_y * y_scale,
            assessment.label,
            color="white",
            fontsize=9,
            fontweight="bold",
            ha="center",
            va="center",
            bbox={
                "facecolor": assessment.color,
                "alpha": 0.9,
                "edgecolor": "white",
                "boxstyle": "round,pad=0.35",
            },
        )

    bar_axis = figure.add_subplot(grid[2, :])
    add_stacked_quality_bar(bar_axis, ripe_pct, unripe_pct, defect_pct)
    return figure


def add_summary_page(document: docx.Document, summaries: list[ImageSummary]) -> None:
    add_heading(document, "Executive Summary", level=1)
    total_images = len(summaries)
    total_tomatoes = sum(summary.tomato_count for summary in summaries)
    total_ripe = sum(summary.ripe_count for summary in summaries)
    total_unripe = sum(summary.unripe_count for summary in summaries)
    total_defective = sum(summary.defective_count for summary in summaries)

    avg_ripe = np.mean([summary.ripe_pct for summary in summaries]) if summaries else 0.0
    avg_unripe = np.mean([summary.unripe_pct for summary in summaries]) if summaries else 0.0
    avg_defect = np.mean([summary.defect_pct for summary in summaries]) if summaries else 0.0

    add_metric_table(
        document,
        [
            ("Images analyzed", str(total_images)),
            ("Total tomatoes / regions", str(total_tomatoes)),
            ("Ripe count", str(total_ripe)),
            ("Unripe count", str(total_unripe)),
            ("Defective count", str(total_defective)),
            ("Average ripe area", f"{avg_ripe:.1f}%"),
            ("Average unripe area", f"{avg_unripe:.1f}%"),
            ("Average defective area", f"{avg_defect:.1f}%"),
        ],
        accent="#E0EAFF",
    )

    document.add_paragraph()
    summary_figure = create_summary_figure(summaries)
    add_figure_to_doc(document, summary_figure, width_inches=7.25)
    plt.close(summary_figure)

    add_heading(document, "Image-Level Results", level=2)
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Image", "Regions", "Ripe", "Unripe", "Defective", "Area Ratio"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, COLORS["navy"])
        set_cell_border(cell, COLORS["navy"])
        set_cell_text(cell, header.upper(), bold=True, color=COLORS["white"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_index, summary in enumerate(summaries, start=1):
        row = table.add_row().cells
        fill = COLORS["white"] if row_index % 2 else COLORS["background"]
        values = [
            summary.image_name,
            str(summary.tomato_count),
            str(summary.ripe_count),
            str(summary.unripe_count),
            str(summary.defective_count),
            f"Ripe {summary.ripe_pct:.1f}% | "
            f"Unripe {summary.unripe_pct:.1f}% | "
            f"Defective {summary.defect_pct:.1f}%",
        ]
        for cell, value in zip(row, values):
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            set_cell_text(cell, value, color=COLORS["navy"], size=8)

    document.add_page_break()


def process_image(
    model: YOLO, image_path: Path, confidence: float, document: docx.Document
) -> ImageSummary:
    print(f"\nAnalyzing: {image_path.name}")
    results = model.predict(source=str(image_path), conf=confidence, verbose=False)
    result = results[0]

    tomatoes, defects, mask_shape = split_detections(result)
    quality = calculate_quality_percentages(tomatoes, defects, mask_shape)

    original_image = Image.open(image_path).convert("RGB")
    image_array = np.array(original_image)

    assessments = assess_tomatoes(tomatoes, defects)
    assessments.extend(assess_standalone_defects(defects, tomatoes, len(assessments)))

    ripe_count = sum(1 for item in assessments if item.status == "Ripe")
    unripe_count = sum(1 for item in assessments if item.status == "Unripe")
    defective_count = sum(1 for item in assessments if item.status == "Defective")

    print(
        "  Result: "
        f"{len(assessments)} regions | "
        f"ripe {quality[0]:.1f}% | unripe {quality[1]:.1f}% | defective {quality[2]:.1f}%"
    )

    add_heading(document, f"Image Analysis: {image_path.name}", level=1)
    add_labeled_paragraph(
        document,
        "Summary",
        f"{len(assessments)} detected region(s), with {defective_count} marked defective by mask overlap analysis.",
    )
    add_metric_table(
        document,
        [
            ("Detected regions", str(len(assessments))),
            ("Ripe", str(ripe_count)),
            ("Unripe", str(unripe_count)),
            ("Defective", str(defective_count)),
            ("Ripe area", f"{quality[0]:.1f}%"),
            ("Unripe area", f"{quality[1]:.1f}%"),
            ("Defective area", f"{quality[2]:.1f}%"),
        ],
        accent="#ECFDF3" if defective_count == 0 else "#FEF3F2",
    )

    plotted_prediction = result.plot()[..., ::-1]
    dashboard = create_dashboard_figure(image_path, original_image, plotted_prediction, assessments, quality, mask_shape)
    add_figure_to_doc(document, dashboard, width_inches=7.3)
    plt.close(dashboard)

    if assessments:
        add_heading(document, "Detected Tomato Details", level=2)
        for assessment in assessments:
            detail = create_detail_figure(assessment, image_array)
            if detail is None:
                continue
            add_figure_to_doc(document, detail, width_inches=7.1)
            plt.close(detail)
    else:
        paragraph = document.add_paragraph()
        paragraph.add_run("No tomato regions were detected in this image.").italic = True

    document.add_section(WD_SECTION.NEW_PAGE)
    return ImageSummary(
        image_name=image_path.name,
        tomato_count=len(assessments),
        ripe_count=ripe_count,
        unripe_count=unripe_count,
        defective_count=defective_count,
        ripe_pct=quality[0],
        unripe_pct=quality[1],
        defect_pct=quality[2],
    )


def main() -> None:
    args = parse_args()
    configure_plotting()
    validate_inputs(args.model, args.images, args.samples)

    image_paths = collect_images(args.images)
    selected_images = select_images(image_paths, args.samples, args.use_random_selection, args.seed)
    report_path = make_report_path(args.output)

    print(f"Report will be saved to: {report_path}")
    print(f"Selected {len(selected_images)} image(s) from {args.images}")
    print(f"Selection mode: {'random' if args.use_random_selection else 'ordered'}")
    if args.use_random_selection:
        print(f"Random seed: {args.seed}")

    model = YOLO(str(args.model))
    document = setup_document()
    add_title_page(document, len(selected_images), args.model, args.images)

    summaries: list[ImageSummary] = []
    try:
        for image_path in selected_images:
            summaries.append(process_image(model, image_path, args.confidence, document))

        add_summary_page(document, summaries)
    finally:
        document.save(report_path)
        print(f"\nProfessional DOCX report saved to: {report_path}")


if __name__ == "__main__":
    main()
