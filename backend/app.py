from __future__ import annotations

import base64
import json
import math
import os
import shutil
from io import BytesIO
from datetime import datetime
from pathlib import Path
from functools import lru_cache
from xml.sax.saxutils import escape as xml_escape

from matplotlib import font_manager as mpl_font_manager
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from PIL import Image, ImageDraw, ImageFont

try:
    from .analysis import analyze_uploaded_image, make_report_payload, report_to_json
except ImportError:  # pragma: no cover - allows running from backend/ directly
    from analysis import analyze_uploaded_image, make_report_payload, report_to_json


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
MODEL_PATH = ROOT_DIR / "best.pt"
DEFAULT_CLASSIFIER_MODEL_PATH = ROOT_DIR / "bestclassifier.pt"
_classifier_path_value = os.environ.get("CLASSIFIER_MODEL_PATH", "").strip()
if _classifier_path_value:
    CLASSIFIER_MODEL_PATH = Path(_classifier_path_value).expanduser()
else:
    CLASSIFIER_MODEL_PATH = DEFAULT_CLASSIFIER_MODEL_PATH if DEFAULT_CLASSIFIER_MODEL_PATH.exists() else None
if CLASSIFIER_MODEL_PATH is not None and not CLASSIFIER_MODEL_PATH.exists():
    CLASSIFIER_MODEL_PATH = DEFAULT_CLASSIFIER_MODEL_PATH if DEFAULT_CLASSIFIER_MODEL_PATH.exists() else None
UPLOAD_DIR = BASE_DIR / "uploads"
REPORT_DIR = BASE_DIR / "reports"

ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
REPORT_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Tomato Grading API", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/api/models")
def models() -> dict:
    return {
        "detector": {
            "path": str(MODEL_PATH),
            "name": MODEL_PATH.name,
            "active": MODEL_PATH.exists(),
        },
        "classifier": {
            "path": str(CLASSIFIER_MODEL_PATH) if CLASSIFIER_MODEL_PATH else None,
            "name": CLASSIFIER_MODEL_PATH.name if CLASSIFIER_MODEL_PATH else None,
            "active": CLASSIFIER_MODEL_PATH is not None,
        },
    }


def load_report(report_id: str) -> dict:
    report_path = REPORT_DIR / f"{report_id}.json"
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found.")
    return json.loads(report_path.read_text(encoding="utf-8"))


def list_reports() -> list[dict]:
    reports: list[dict] = []
    for report_path in sorted(REPORT_DIR.glob("*.json"), reverse=True):
        try:
            report = json.loads(report_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            continue
        reports.append(
            {
                "reportId": report.get("reportId", report_path.stem),
                "generatedAt": report.get("generatedAt"),
                "fileName": report.get("fileName"),
                "summary": report.get("summary"),
                "totals": report.get("totals", {}),
                "tomatoCount": report.get("tomatoCount", 0),
                "ripeCount": report.get("ripeCount", 0),
                "unripeCount": report.get("unripeCount", 0),
                "defectiveCount": report.get("defectiveCount", 0),
                "ripePct": report.get("ripePct", 0),
                "unripePct": report.get("unripePct", 0),
                "defectPct": report.get("defectPct", 0),
                "modelInfo": report.get("modelInfo", {}),
                "classifierUsedCount": report.get("classifierUsedCount", 0),
            }
        )
    return reports


def format_datetime(value: str | None) -> str:
    if not value:
        return "Unknown"
    try:
        parsed = datetime.fromisoformat(value)
    except ValueError:
        return value
    return parsed.strftime("%Y-%m-%d %H:%M:%S")


def shorten_text(value: str | None, max_length: int = 40) -> str:
    if not value:
        return "Unknown"
    if len(value) <= max_length:
        return value
    head = max(8, math.floor(max_length * 0.6))
    tail = max(6, max_length - head - 1)
    return f"{value[:head]}…{value[-tail:]}"


def shorten_text_clean(value: str | None, max_length: int = 40) -> str:
    if not value:
        return "Unknown"
    if len(value) <= max_length:
        return value
    head = max(8, math.floor(max_length * 0.6))
    tail = max(6, max_length - head - 3)
    return f"{value[:head]}...{value[-tail:]}"


def shorten_text(value: str | None, max_length: int = 40) -> str:
    if not value:
        return "Unknown"
    if len(value) <= max_length:
        return value
    head = max(8, math.floor(max_length * 0.6))
    tail = max(6, max_length - head - 3)
    return f"{value[:head]}...{value[-tail:]}"


def build_report_sections(report: dict) -> list[str]:
    summary_lines = [
        f"Report ID: {report.get('reportId', 'Unknown')}",
        f"Generated At: {format_datetime(report.get('generatedAt'))}",
        f"Source File: {report.get('fileName', 'Unknown')}",
        "",
        "Executive Summary",
        report.get("summary", ""),
        "",
        "Quality Metrics",
        f"Regions detected: {report.get('tomatoCount', 0)}",
        f"Ripe: {report.get('ripeCount', 0)} ({report.get('ripePct', 0):.1f}%)",
        f"Unripe: {report.get('unripeCount', 0)} ({report.get('unripePct', 0):.1f}%)",
        f"Defective: {report.get('defectiveCount', 0)} ({report.get('defectPct', 0):.1f}%)",
        "",
        "Detected Regions",
    ]

    for item in report.get("assessments", []):
        summary_lines.append(
            f"{item.get('number', 0):>2}. {item.get('label', 'Unknown')} | "
            f"Status: {item.get('status', 'Unknown')} | "
            f"Confidence: {float(item.get('confidence', 0)):0.3f} | "
            f"Defect: {float(item.get('defectPercent', 0)):0.1f}%"
        )

    if not report.get("assessments"):
        summary_lines.append("No regions were detected for this image.")

    return summary_lines


PDF_PAGE_SIZE = (1240, 1754)
PDF_MARGIN = 70
PDF_BG = (247, 243, 236)
PDF_PANEL = (255, 255, 255)
PDF_TEXT = (17, 30, 44)
PDF_MUTED = (102, 112, 133)
PDF_LINE = (215, 221, 230)
PDF_ACCENT = (217, 119, 6)
PDF_GREEN = (46, 125, 50)
PDF_AMBER = (185, 119, 0)
PDF_RED = (198, 40, 40)


@lru_cache(maxsize=16)
def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = mpl_font_manager.findfont("DejaVu Sans Bold" if bold else "DejaVu Sans")
    return ImageFont.truetype(font_path, size=size)


def decode_data_uri(data_uri: str | None) -> Image.Image | None:
    if not data_uri:
        return None
    try:
        _, payload = data_uri.split(",", 1)
    except ValueError:
        payload = data_uri
    try:
        return Image.open(BytesIO(base64.b64decode(payload))).convert("RGB")
    except Exception:  # noqa: BLE001
        return None


def image_with_padding(image: Image.Image, size: tuple[int, int], fill: tuple[int, int, int] = (255, 255, 255)) -> Image.Image:
    target_w, target_h = size
    src_w, src_h = image.size
    scale = min(target_w / src_w, target_h / src_h)
    resized = image.resize((max(1, int(src_w * scale)), max(1, int(src_h * scale))), Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", size, fill)
    offset = ((target_w - resized.width) // 2, (target_h - resized.height) // 2)
    canvas.paste(resized, offset)
    return canvas


def text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> int:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0]


def text_height(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> int:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    for paragraph in str(text).splitlines():
        if not paragraph.strip():
            lines.append("")
            continue
        current = ""
        for word in paragraph.split():
            candidate = word if not current else f"{current} {word}"
            if text_width(draw, candidate, font) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines or [""]


def draw_text_block(
    draw: ImageDraw.ImageDraw,
    x: int,
    y: int,
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
    max_width: int,
    line_gap: int = 8,
) -> int:
    for line in wrap_text(draw, text, font, max_width):
        if line:
            draw.text((x, y), line, font=font, fill=fill)
            y += text_height(draw, line, font) + line_gap
        else:
            y += text_height(draw, "Ag", font) // 2
    return y


def draw_rounded_image(
    base: Image.Image,
    image: Image.Image,
    box: tuple[int, int, int, int],
    radius: int = 28,
) -> None:
    x1, y1, x2, y2 = box
    target = image_with_padding(image, (x2 - x1, y2 - y1), fill=PDF_PANEL)
    mask = Image.new("L", target.size, 0)
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle((0, 0, target.width - 1, target.height - 1), radius=radius, fill=255)
    base.paste(target, (x1, y1), mask)


def draw_header(draw: ImageDraw.ImageDraw, report: dict, page_width: int) -> int:
    title_font = load_font(34, bold=True)
    body_font = load_font(18)
    label_font = load_font(13, bold=True)
    x = PDF_MARGIN
    y = PDF_MARGIN
    draw.text((x, y), "Tomato Grading Report", font=title_font, fill=PDF_TEXT)
    y += text_height(draw, "Ag", title_font) + 8
    draw.text((x, y), shorten_text_clean(report.get("fileName"), 56), font=body_font, fill=PDF_MUTED)
    y += text_height(draw, "Ag", body_font) + 16
    meta = f"Report ID: {report.get('reportId', 'Unknown')}   |   Generated: {format_datetime(report.get('generatedAt'))}"
    draw.text((x, y), meta, font=label_font, fill=PDF_ACCENT)
    pill_text = "Ready"
    pill_font = load_font(16, bold=True)
    pill_w = text_width(draw, pill_text, pill_font) + 36
    pill_h = 42
    pill_x2 = page_width - PDF_MARGIN
    pill_x1 = pill_x2 - pill_w
    pill_y1 = PDF_MARGIN + 2
    pill_y2 = pill_y1 + pill_h
    draw.rounded_rectangle((pill_x1, pill_y1, pill_x2, pill_y2), radius=21, fill=(228, 238, 224), outline=None)
    draw.text((pill_x1 + (pill_w - text_width(draw, pill_text, pill_font)) / 2, pill_y1 + 11), pill_text, font=pill_font, fill=PDF_GREEN)
    return y + 34


def draw_section_card(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, title_font, body_font) -> None:
    draw.rounded_rectangle(box, radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
    x1, y1, _, _ = box
    draw.text((x1 + 22, y1 + 18), title, font=title_font, fill=PDF_TEXT)


def create_overview_page(report: dict) -> Image.Image:
    page = Image.new("RGB", PDF_PAGE_SIZE, PDF_BG)
    draw = ImageDraw.Draw(page)
    title_font = load_font(24, bold=True)
    small_font = load_font(15)
    body_font = load_font(18)
    heading_font = load_font(18, bold=True)
    number_font = load_font(32, bold=True)

    content_top = draw_header(draw, report, PDF_PAGE_SIZE[0])

    # Summary card
    summary_box = (PDF_MARGIN, content_top, PDF_PAGE_SIZE[0] - PDF_MARGIN, content_top + 148)
    draw_section_card(draw, summary_box, "Summary", heading_font, body_font)
    summary_text = report.get("summary", "No summary available.")
    y = draw_text_block(draw, summary_box[0] + 22, summary_box[1] + 56, summary_text, body_font, PDF_TEXT, summary_box[2] - summary_box[0] - 44, 7)

    # Metric cards
    metric_top = summary_box[3] + 18
    metric_h = 120
    metric_gap = 16
    metric_w = (PDF_PAGE_SIZE[0] - PDF_MARGIN * 2 - metric_gap * 3) // 4
    metrics = [
        ("Regions", report.get("tomatoCount", 0), PDF_TEXT),
        ("Ripe", report.get("ripeCount", 0), PDF_GREEN),
        ("Unripe", report.get("unripeCount", 0), PDF_AMBER),
        ("Defective", report.get("defectiveCount", 0), PDF_RED),
    ]
    for index, (label, value, color) in enumerate(metrics):
        x1 = PDF_MARGIN + index * (metric_w + metric_gap)
        box = (x1, metric_top, x1 + metric_w, metric_top + metric_h)
        draw.rounded_rectangle(box, radius=24, fill=PDF_PANEL, outline=PDF_LINE, width=1)
        draw.text((box[0] + 18, box[1] + 18), label.upper(), font=small_font, fill=PDF_MUTED)
        draw.text((box[0] + 18, box[1] + 54), str(value), font=number_font, fill=color)

    # Images section
    images_top = metric_top + metric_h + 18
    image_box_h = 430
    image_box_w = (PDF_PAGE_SIZE[0] - PDF_MARGIN * 2 - 18) // 2
    image_titles = [("Original", decode_data_uri(report.get("image"))), ("Annotated", decode_data_uri(report.get("annotatedImage")))]
    for index, (title, image) in enumerate(image_titles):
        x1 = PDF_MARGIN + index * (image_box_w + 18)
        box = (x1, images_top, x1 + image_box_w, images_top + image_box_h)
        draw.rounded_rectangle(box, radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
        draw.text((box[0] + 22, box[1] + 18), title, font=heading_font, fill=PDF_TEXT)
        if image is not None:
            draw_rounded_image(page, image, (box[0] + 12, box[1] + 56, box[2] - 12, box[3] - 12), radius=22)

    # Chart below, if present
    chart = decode_data_uri(report.get("summaryChart"))
    chart_top = images_top + image_box_h + 18
    chart_box = (PDF_MARGIN, chart_top, PDF_PAGE_SIZE[0] - PDF_MARGIN, chart_top + 260)
    draw.rounded_rectangle(chart_box, radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
    draw.text((chart_box[0] + 22, chart_box[1] + 18), "Chart", font=heading_font, fill=PDF_TEXT)
    if chart is not None:
        draw_rounded_image(page, chart, (chart_box[0] + 12, chart_box[1] + 56, chart_box[2] - 12, chart_box[3] - 12), radius=22)

    # Footer note
    footer_font = load_font(13)
    footer_text = "Generated from the selected image or webcam frame."
    draw.text((PDF_MARGIN, PDF_PAGE_SIZE[1] - PDF_MARGIN - 18), footer_text, font=footer_font, fill=PDF_MUTED)
    return page


def create_assessment_pages(report: dict) -> list[Image.Image]:
    pages: list[Image.Image] = []
    rows = report.get("assessments", [])
    if not rows:
        page = Image.new("RGB", PDF_PAGE_SIZE, PDF_BG)
        draw = ImageDraw.Draw(page)
        draw_header(draw, report, PDF_PAGE_SIZE[0])
        draw.rounded_rectangle((PDF_MARGIN, 260, PDF_PAGE_SIZE[0] - PDF_MARGIN, 520), radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
        draw.text((PDF_MARGIN + 24, 290), "Detected Regions", font=load_font(22, bold=True), fill=PDF_TEXT)
        draw.text((PDF_MARGIN + 24, 340), "No regions were detected for this image.", font=load_font(18), fill=PDF_MUTED)
        pages.append(page)
        return pages

    row_font = load_font(16)
    head_font = load_font(16, bold=True)
    title_font = load_font(24, bold=True)
    for start in range(0, len(rows), 18):
        page = Image.new("RGB", PDF_PAGE_SIZE, PDF_BG)
        draw = ImageDraw.Draw(page)
        draw_header(draw, report, PDF_PAGE_SIZE[0])
        box = (PDF_MARGIN, 220, PDF_PAGE_SIZE[0] - PDF_MARGIN, PDF_PAGE_SIZE[1] - PDF_MARGIN)
        draw.rounded_rectangle(box, radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
        draw.text((box[0] + 22, box[1] + 18), "Detected Regions", font=title_font, fill=PDF_TEXT)

        table_x = box[0] + 20
        table_y = box[1] + 66
        cols = [64, 180, 200, 160, 120]
        headers = ["#", "Status", "Label", "Confidence", "Defect"]
        row_h = 34
        header_h = 34
        draw.rounded_rectangle((table_x, table_y, table_x + sum(cols), table_y + header_h), radius=12, fill=(241, 245, 249), outline=PDF_LINE, width=1)
        cx = table_x + 12
        for idx, header in enumerate(headers):
            draw.text((cx, table_y + 9), header, font=head_font, fill=PDF_MUTED)
            cx += cols[idx]

        y = table_y + header_h + 8
        for row in rows[start:start + 18]:
            if y + row_h > box[3] - 20:
                break
            draw.rounded_rectangle((table_x, y, table_x + sum(cols), y + row_h), radius=10, fill=PDF_PANEL if (row["number"] % 2) else PDF_PANEL, outline=PDF_LINE, width=1)
            cells = [
                str(row.get("number", "")),
                row.get("status", ""),
                shorten_text_clean(row.get("label"), 26),
                f"{float(row.get('confidence', 0)):0.3f}",
                f"{float(row.get('defectPercent', 0)):0.1f}%",
            ]
            cx = table_x + 12
            for idx, cell in enumerate(cells):
                draw.text((cx, y + 8), cell, font=row_font, fill=PDF_TEXT)
                cx += cols[idx]
            y += row_h + 6
        pages.append(page)
    return pages


def create_detail_pages(report: dict) -> list[Image.Image]:
    detail_images = report.get("detailImages", [])
    if not detail_images:
        return []

    pages: list[Image.Image] = []
    title_font = load_font(24, bold=True)
    heading_font = load_font(18, bold=True)
    small_font = load_font(14)
    for start in range(0, len(detail_images), 4):
        page = Image.new("RGB", PDF_PAGE_SIZE, PDF_BG)
        draw = ImageDraw.Draw(page)
        draw_header(draw, report, PDF_PAGE_SIZE[0])
        draw.rounded_rectangle((PDF_MARGIN, 220, PDF_PAGE_SIZE[0] - PDF_MARGIN, PDF_PAGE_SIZE[1] - PDF_MARGIN), radius=28, fill=PDF_PANEL, outline=PDF_LINE, width=1)
        draw.text((PDF_MARGIN + 22, 238), "Region Details", font=title_font, fill=PDF_TEXT)
        card_w = (PDF_PAGE_SIZE[0] - PDF_MARGIN * 2 - 18) // 2
        card_h = 300
        for idx, item in enumerate(detail_images[start:start + 4]):
            col = idx % 2
            row = idx // 2
            x1 = PDF_MARGIN + col * (card_w + 18)
            y1 = 286 + row * (card_h + 18)
            box = (x1, y1, x1 + card_w, y1 + card_h)
            draw.rounded_rectangle(box, radius=22, fill=PDF_PANEL, outline=PDF_LINE, width=1)
            draw.text((box[0] + 18, box[1] + 16), item.get("title", "Detail"), font=heading_font, fill=PDF_TEXT)
            draw.text((box[0] + 18, box[1] + 42), item.get("caption", ""), font=small_font, fill=PDF_MUTED)
            img = decode_data_uri(item.get("image"))
            if img is not None:
                draw_rounded_image(page, img, (box[0] + 12, box[1] + 70, box[2] - 12, box[3] - 12), radius=18)
        pages.append(page)
    return pages


def build_pdf_bytes(report: dict) -> bytes:
    pages = [create_overview_page(report), *create_assessment_pages(report), *create_detail_pages(report)]
    if not pages:
        pages = [Image.new("RGB", PDF_PAGE_SIZE, PDF_BG)]
    buffer = BytesIO()
    pages[0].save(buffer, format="PDF", save_all=True, append_images=pages[1:])
    return buffer.getvalue()


def _docx_run(text: str, *, bold: bool = False, size: int = 22) -> str:
    props = [f'<w:sz w:val="{size}"/>', f'<w:szCs w:val="{size}"/>']
    if bold:
        props.insert(0, "<w:b/>")
    return f"<w:r><w:rPr>{''.join(props)}</w:rPr><w:t xml:space='preserve'>{xml_escape(text)}</w:t></w:r>"


def _docx_paragraph(text: str, *, bold: bool = False, size: int = 22, align: str | None = None) -> str:
    ppr = []
    if align:
        ppr.append(f'<w:jc w:val="{align}"/>')
    return f"<w:p><w:pPr>{''.join(ppr)}</w:pPr>{_docx_run(text, bold=bold, size=size)}</w:p>"


def build_docx_bytes(report: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tomato Grading Report")
    run.bold = True
    run.font.size = Pt(20)

    meta = document.add_paragraph()
    meta.add_run(f"Report ID: {report.get('reportId', 'Unknown')}\n").bold = True
    meta.add_run(f"Generated At: {format_datetime(report.get('generatedAt'))}\n")
    meta.add_run(f"Source File: {report.get('fileName', 'Unknown')}")

    summary_heading = document.add_paragraph()
    summary_heading.add_run("Executive Summary").bold = True
    document.add_paragraph(report.get("summary", "No summary available."))

    metrics_heading = document.add_paragraph()
    metrics_heading.add_run("Quality Metrics").bold = True
    metrics = document.add_table(rows=1, cols=2)
    metrics.style = "Table Grid"
    metric_header = metrics.rows[0].cells
    metric_header[0].text = "Metric"
    metric_header[1].text = "Value"
    for label, value in (
        ("Regions detected", report.get("tomatoCount", 0)),
        ("Ripe", f"{report.get('ripeCount', 0)} ({report.get('ripePct', 0):.1f}%)"),
        ("Unripe", f"{report.get('unripeCount', 0)} ({report.get('unripePct', 0):.1f}%)"),
        ("Defective", f"{report.get('defectiveCount', 0)} ({report.get('defectPct', 0):.1f}%)"),
    ):
        row = metrics.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    image_sections = [
        ("Original Image", report.get("image")),
        ("Annotated Image", report.get("annotatedImage")),
        ("Summary Chart", report.get("summaryChart")),
    ]

    for title_text, image_data in image_sections:
        img = decode_data_uri(image_data)
        if img is None:
            continue
        document.add_paragraph().add_run(title_text).bold = True
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        document.add_picture(buffer, width=Inches(6.5))

    regions_heading = document.add_paragraph()
    regions_heading.add_run("Detected Regions").bold = True
    if report.get("assessments"):
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "#"
        headers[1].text = "Status"
        headers[2].text = "Label"
        headers[3].text = "Confidence / Defect"
        for item in report["assessments"]:
            row = table.add_row().cells
            row[0].text = str(item.get("number", 0))
            row[1].text = str(item.get("status", "Unknown"))
            row[2].text = str(item.get("label", "Unknown"))
            row[3].text = f"{float(item.get('confidence', 0)):0.3f} / {float(item.get('defectPercent', 0)):0.1f}%"
    else:
        document.add_paragraph("No regions were detected for this image.")

    if report.get("detailImages"):
        details_heading = document.add_paragraph()
        details_heading.add_run("Region Details").bold = True
        for item in report["detailImages"]:
            caption = f"{item.get('title', 'Detail')}: {item.get('caption', '')}"
            document.add_paragraph(caption)
            img = decode_data_uri(item.get("image"))
            if img is None:
                continue
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            buffer.seek(0)
            document.add_picture(buffer, width=Inches(6.5))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@app.post("/api/analyze")
async def analyze(file: UploadFile = File(...), confidence: float = 0.25) -> dict:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload a JPG, PNG, BMP, or WEBP image.")

    report_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    stored_name = f"{report_id}_{Path(file.filename).name}"
    stored_path = UPLOAD_DIR / stored_name

    with stored_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        report = analyze_uploaded_image(
            MODEL_PATH,
            stored_path,
            confidence=confidence,
            classifier_path=CLASSIFIER_MODEL_PATH,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Analysis failed: {exc}") from exc

    payload = make_report_payload(report)
    payload["reportId"] = report_id
    payload["storedImage"] = stored_name
    payload["modelInfo"] = {
        "detector": MODEL_PATH.name,
        "classifier": CLASSIFIER_MODEL_PATH.name if CLASSIFIER_MODEL_PATH else None,
        "classifierActive": CLASSIFIER_MODEL_PATH is not None,
    }
    payload["classifierUsedCount"] = sum(
        1 for item in payload.get("assessments", []) if item.get("classifierLabel")
    )

    report_path = REPORT_DIR / f"{report_id}.json"
    report_path.write_text(json.dumps(report_to_json(payload), indent=2), encoding="utf-8")

    return payload


@app.get("/api/reports/{report_id}")
def get_report(report_id: str) -> dict:
    return load_report(report_id)


@app.get("/api/reports")
def get_reports() -> dict:
    return {"reports": list_reports()}


@app.get("/api/reports/{report_id}/download")
def download_report(report_id: str, format: str = "json") -> Response:
    report = load_report(report_id)
    download_format = format.lower()

    if download_format == "json":
        payload = json.dumps(report, indent=2).encode("utf-8")
        return Response(
            content=payload,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="tomato-report-{report_id}.json"'},
        )

    if download_format == "pdf":
        payload = build_pdf_bytes(report)
        return Response(
            content=payload,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="tomato-report-{report_id}.pdf"'},
        )

    if download_format in {"doc", "docx"}:
        payload = build_docx_bytes(report)
        return Response(
            content=payload,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="tomato-report-{report_id}.docx"'},
        )

    raise HTTPException(status_code=400, detail="Unsupported export format. Use json, pdf, or docx.")
